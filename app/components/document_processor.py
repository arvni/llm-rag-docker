import io
import os
import tempfile
import logging
from typing import List, Optional, Union
from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import cv2
import numpy as np
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import streamlit as st

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced document processor with OCR capabilities"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._process_pdf,
            'txt': self._process_txt,
            'docx': self._process_docx,
            'png': self._process_image,
            'jpg': self._process_image,
            'jpeg': self._process_image
        }
        
        # OCR configuration
        self.ocr_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789 .,!?;:()'
        
    def process_file(self, uploaded_file, ocr_enabled: bool = True, 
                    chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
        """Process uploaded file and return document chunks"""
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            # Extract text based on file type
            processor = self.supported_formats[file_extension]
            text_content = processor(uploaded_file, ocr_enabled)
            
            if not text_content.strip():
                raise ValueError("No text content extracted from file")
            
            # Create text splitter
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
            
            # Split text into chunks
            chunks = text_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        'filename': uploaded_file.name,
                        'file_type': file_extension,
                        'chunk_id': i,
                        'total_chunks': len(chunks),
                        'ocr_enabled': ocr_enabled
                    }
                )
                documents.append(doc)
            
            logger.info(f"Processed {uploaded_file.name}: {len(chunks)} chunks created")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing file {uploaded_file.name}: {e}")
            raise
    
    def _process_pdf(self, uploaded_file, ocr_enabled: bool = True) -> str:
        """Process PDF file with OCR fallback"""
        text_content = []
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name
        
        try:
            doc = fitz.open(tmp_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                # If no text found and OCR is enabled, use OCR
                if not text.strip() and ocr_enabled:
                    st.info(f"Using OCR for page {page_num + 1}")
                    
                    # Convert page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Higher resolution
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))
                    
                    # Enhance image for better OCR
                    text = self._ocr_image(image)
                
                if text.strip():
                    text_content.append(f"Page {page_num + 1}:\n{text}")
            
            doc.close()
            
        finally:
            os.unlink(tmp_path)
        
        return "\n\n".join(text_content)
    
    def _process_txt(self, uploaded_file, ocr_enabled: bool = False) -> str:
        """Process plain text file"""
        try:
            content = uploaded_file.read().decode('utf-8')
            return content
        except UnicodeDecodeError:
            # Try other encodings
            uploaded_file.seek(0)
            content = uploaded_file.read().decode('latin-1')
            return content
    
    def _process_docx(self, uploaded_file, ocr_enabled: bool = False) -> str:
        """Process DOCX file"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name
        
        try:
            doc = docx.Document(tmp_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        finally:
            os.unlink(tmp_path)
    
    def _process_image(self, uploaded_file, ocr_enabled: bool = True) -> str:
        """Process image file with OCR"""
        if not ocr_enabled:
            return ""
        
        image = Image.open(uploaded_file)
        return self._ocr_image(image)
    
    def _ocr_image(self, image: Image.Image) -> str:
        """Perform OCR on image with preprocessing"""
        try:
            # Convert to numpy array
            img_array = np.array(image)
            
            # Convert to grayscale if needed
            if len(img_array.shape) == 3:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            
            # Image preprocessing for better OCR
            img_array = self._preprocess_image(img_array)
            
            # Perform OCR
            text = pytesseract.image_to_string(
                Image.fromarray(img_array), 
                config=self.ocr_config
            )
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"OCR error: {e}")
            return ""
    
    def _preprocess_image(self, img_array: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR results"""
        # Noise removal
        img_array = cv2.medianBlur(img_array, 5)
        
        # Thresholding
        img_array = cv2.threshold(img_array, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
        
        # Morphological operations
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
        img_array = cv2.morphologyEx(img_array, cv2.MORPH_CLOSE, kernel)
        
        return img_array
    
    def get_file_info(self, uploaded_file) -> dict:
        """Get file information"""
        return {
            'name': uploaded_file.name,
            'size': len(uploaded_file.read()) if hasattr(uploaded_file, 'read') else 0,
            'type': uploaded_file.type if hasattr(uploaded_file, 'type') else 'unknown'
        }